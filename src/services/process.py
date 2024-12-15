from PyPDF2 import PdfReader
from docx import Document
import openai
from nltk.tokenize import sent_tokenize
import os
import time
import json
import re
from collections import defaultdict
from docx import Document

import nltk
nltk.download('punkt')


openai.api_key = os.getenv("OPENAI_API_KEY")

def extract_text(file):
    """
    Args:
        file: Input File Uploaded

    Returns:
        Contents of the file
    """
    if file.type == "application/pdf":
        pdf_reader = PdfReader(file)
        text = " ".join(page.extract_text() for page in pdf_reader.pages)
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(file)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    else:
        text = file.read().decode("utf-8")
    print(text)
    return text


def chunk_text(text, max_tokens=3000):
    """
    Args:
        text: Contents fetched from the input file
        max_tokens: Max token at once

    Returns:
        Chunks of texts
    """
    sentences = sent_tokenize(text)
    chunks = []
    current_chunk = []
    current_length = 0

    for sentence in sentences:
        sentence_length = len(sentence.split())
        if current_length + sentence_length > max_tokens:
            chunks.append(" ".join(current_chunk))
            current_chunk = []
            current_length = 0

        current_chunk.append(sentence)
        current_length += sentence_length

    if current_chunk:
        chunks.append(" ".join(current_chunk))

    return chunks

# def generate_user_stories_and_criteria(text):
#     """
#     Args:
#         text: Contents of the input file

#     Returns:
#         Generated user stories and acceptance criteria from the model.
#     """
#     chunks = chunk_text(text)
#     user_stories = []


#     for chunk in chunks:
#         messages = [
#             {"role": "system", "content": "You are a helpful assistant."},
#             {"role": "user", "content": f'''Based on the following requirements:
#              {chunk}
#             Generate user stories and acceptance criteria for each requirement.'''}
#         ]
#         response = openai.ChatCompletion.create(
#         model="gpt-4o",
#         messages=messages,
#         max_tokens=1500,
#         temperature=0
#         )
#         user_stories.append(response.choices[0].message['content'].strip())

#     return "".join(user_stories)

def generate_user_stories_and_criteria_specific(text):
    """
    Args:
        text: Contents of the input file

    Returns:
        Generated user stories and acceptance criteria from the model.
    """
    chunks = chunk_text(text)
    user_stories = []
    for chunk in chunks:
        prompt = f"""
                You are tasked with creating User Stories and Acceptance Criteria from the given document. 
                {chunk}
                Follow these instructions:
                
    
                1. Identify all **Modules** in the document (e.g., Recruitment & Onboarding, Attendance, Leave Management, etc.).
                2. For each module, extract individual **User Stories** based on the functionalities described.
                3. For each User Story:
                   - Assign a **User Story Title** that reflects the functionality.
                   - Define the **Roles** involved (e.g., HR, Recruiter, Employee, Manager).
                   - List the **Actions** that each role can perform under the User Story.
                   - For each action, specify the **arguments** required (e.g., fields like title, description, date).
                
                Use the following format for each module:
                
                Module Name: [Module Name]
                User Story [X]: [User Story Title]
                Acceptance Criteria:
                Role: [Role Name]
                - [Action 1]
                  arguments: [Argument 1, Argument 2, ...]
                - [Action 2]
                  arguments: [Argument 1, Argument 2, ...]
                
                Role: [Role Name]
                - [Action 1]
                  arguments: [Argument 1, Argument 2, ...]
                - [Action 2]
                  arguments: [Argument 1, Argument 2, ...]
                
                Ensure the output is structured and consistent. Generate one User Story for each identified functionality under the module, clearly distinguishing between roles and their actions.
                
                Example:
                
                Module Name: Recruitment & Onboarding
                User Story 1: Job Postings
                Acceptance Criteria:
                Role: Recruiter
                - create job
                  arguments: title, description, requirements
                - edit job
                  arguments: title, description, requirements
                
                Role: Candidate
                - view jobs
                  arguments: filters (location, role, etc.)
                - apply for jobs
                  arguments: resume, cover letter

        """
        messages = [
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt.strip()}
        ]
        response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=messages,
        max_tokens=1500,
        temperature=0
        )
        user_stories.append(response.choices[0].message['content'].strip())

    return "".join(user_stories)

def compare_dom_tech_doc(dom_structure=[], tech_doc={}):
    dom_structure = [
            {
                "lang": "en",
                "class": "hydrated"
            },
            {},
            {
                "charset": "UTF-8"
            },
            {
                "data-styles": ""
            },
            {
                "http-equiv": "X-UA-Compatible",
                "content": "IE=edge"
            },
            {
                "name": "viewport",
                "content": "width=device-width, initial-scale=1.0"
            },
            {},
            {
                "rel": "apple-touch-icon",
                "sizes": "180x180",
                "href": "/static/favicons/apple-touch-icon.png"
            },
            {
                "rel": "icon",
                "type": "image/png",
                "sizes": "32x32",
                "href": "/static/favicons/favicon-32x32.png"
            },
            {
                "rel": "icon",
                "type": "image/png",
                "sizes": "16x16",
                "href": "/static/favicons/favicon-16x16.png"
            },
            {
                "rel": "stylesheet",
                "href": "/static/build/css/style.min.css"
            },
            {
                "rel": "manifest",
                "href": "/static/build/manifest.json"
            },
            {},
            {},
            {},
            {},
            {},
            {},
            {},
            {
                "id": "main"
            },
            {
                "class": "oh-alert-container"
            },
            {
                "class": "oh-auth"
            },
            {
                "class": "oh-auth-card mb-4"
            },
            {
                "class": "oh-onboarding-card__title oh-onboarding-card__title--h2 text-center my-3"
            },
            {
                "class": "text-muted text-center"
            },
            {
                "method": "post",
                "class": "oh-form-group"
            },
            {
                "type": "hidden",
                "name": "csrfmiddlewaretoken",
                "value": "8wFOC361t5dRuE8rTN5GVQK1SZ601tUMkniWhFNrK1KUXtXuM0BjV0PCUujFPSuY"
            },
            {
                "class": "oh-input-group"
            },
            {
                "class": "oh-label",
                "for": "username"
            },
            {
                "type": "text",
                "id": "username",
                "name": "username",
                "class": "oh-input w-100",
                "placeholder": "e.g. adam.luis@horilla.com"
            },
            {
                "class": "oh-input-group"
            },
            {
                "class": "oh-label",
                "for": "password"
            },
            {
                "class": "oh-password-input-container"
            },
            {
                "type": "password",
                "id": "password",
                "name": "password",
                "class": "oh-input oh-input--password w-100",
                "placeholder": "Use alphanumeric characters"
            },
            {
                "type": "button",
                "class": "oh-btn oh-btn--transparent oh-password-input--toggle"
            },
            {
                "class": "oh-passowrd-input__show-icon md hydrated",
                "title": "Show Password",
                "name": "eye-outline",
                "role": "img",
                "aria-label": "eye outline"
            },
            {
                "class": "oh-passowrd-input__hide-icon d-none md hydrated",
                "title": "Hide Password",
                "name": "eye-off-outline",
                "role": "img",
                "aria-label": "eye off outline"
            },
            {
                "type": "submit",
                "class": "oh-btn oh-onboarding-card__button mt-4 oh-btn--secondary oh-btn--shadow w-100 mb-4",
                "role": "button"
            },
            {
                "class": "me-2 md hydrated",
                "name": "lock-closed-outline",
                "role": "img",
                "aria-label": "lock closed outline"
            },
            {
                "class": "text-center"
            },
            {
                "href": "/forgot-password",
                "class": "oh-link oh-link--secondary justify-content-center"
            },
            {},
            {
                "src": "/static/images/ui/auth-logo.png",
                "alt": "Horilla"
            },
            {
                "src": "/static/build/js/web.frontend.min.js"
            },
            {
                "role": "log",
                "aria-live": "assertive",
                "aria-relevant": "additions",
                "class": "ui-helper-hidden-accessible"
            },
            {
                "src": "https://code.jquery.com/jquery-3.6.4.min.js"
            },
            {
                "type": "module",
                "src": "https://unpkg.com/ionicons@5.5.2/dist/ionicons/ionicons.esm.js"
            },
            {
                "nomodule": "",
                "src": "https://unpkg.com/ionicons@5.5.2/dist/ionicons/ionicons.js"
            },
            {
                "src": "https://cdn.jsdelivr.net/npm/sweetalert2@10"
            },
            {},
            {
                "class": "icon-inner"
            },
            {
                "xmlns": "http://www.w3.org/2000/svg",
                "class": "ionicon s-ion-icon",
                "viewBox": "0 0 512 512"
            },
            {},
            {
                "d": "M255.66 112c-77.94 0-157.89 45.11-220.83 135.33a16 16 0 00-.27 17.77C82.92 340.8 161.8 400 255.66 400c92.84 0 173.34-59.38 221.79-135.25a16.14 16.14 0 000-17.47C428.89 172.28 347.8 112 255.66 112z",
                "stroke-linecap": "round",
                "stroke-linejoin": "round",
                "class": "ionicon-fill-none ionicon-stroke-width"
            },
            {
                "cx": "256",
                "cy": "256",
                "r": "80",
                "stroke-miterlimit": "10",
                "class": "ionicon-fill-none ionicon-stroke-width"
            },
            {
                "class": "icon-inner"
            },
            {
                "xmlns": "http://www.w3.org/2000/svg",
                "class": "ionicon s-ion-icon",
                "viewBox": "0 0 512 512"
            },
            {},
            {
                "d": "M432 448a15.92 15.92 0 01-11.31-4.69l-352-352a16 16 0 0122.62-22.62l352 352A16 16 0 01432 448zM255.66 384c-41.49 0-81.5-12.28-118.92-36.5-34.07-22-64.74-53.51-88.7-91v-.08c19.94-28.57 41.78-52.73 65.24-72.21a2 2 0 00.14-2.94L93.5 161.38a2 2 0 00-2.71-.12c-24.92 21-48.05 46.76-69.08 76.92a31.92 31.92 0 00-.64 35.54c26.41 41.33 60.4 76.14 98.28 100.65C162 402 207.9 416 255.66 416a239.13 239.13 0 0075.8-12.58 2 2 0 00.77-3.31l-21.58-21.58a4 4 0 00-3.83-1 204.8 204.8 0 01-51.16 6.47zM490.84 238.6c-26.46-40.92-60.79-75.68-99.27-100.53C349 110.55 302 96 255.66 96a227.34 227.34 0 00-74.89 12.83 2 2 0 00-.75 3.31l21.55 21.55a4 4 0 003.88 1 192.82 192.82 0 0150.21-6.69c40.69 0 80.58 12.43 118.55 37 34.71 22.4 65.74 53.88 89.76 91a.13.13 0 010 .16 310.72 310.72 0 01-64.12 72.73 2 2 0 00-.15 2.95l19.9 19.89a2 2 0 002.7.13 343.49 343.49 0 0068.64-78.48 32.2 32.2 0 00-.1-34.78z"
            },
            {
                "d": "M256 160a95.88 95.88 0 00-21.37 2.4 2 2 0 00-1 3.38l112.59 112.56a2 2 0 003.38-1A96 96 0 00256 160zM165.78 233.66a2 2 0 00-3.38 1 96 96 0 00115 115 2 2 0 001-3.38z"
            },
            {
                "class": "icon-inner"
            },
            {
                "xmlns": "http://www.w3.org/2000/svg",
                "class": "ionicon s-ion-icon",
                "viewBox": "0 0 512 512"
            },
            {},
            {
                "d": "M336 208v-95a80 80 0 00-160 0v95",
                "stroke-linecap": "round",
                "stroke-linejoin": "round",
                "class": "ionicon-fill-none ionicon-stroke-width"
            },
            {
                "x": "96",
                "y": "208",
                "width": "320",
                "height": "272",
                "rx": "48",
                "ry": "48",
                "stroke-linecap": "round",
                "stroke-linejoin": "round",
                "class": "ionicon-fill-none ionicon-stroke-width"
            }
        ]

    tech_doc = {
        "Module Name: Login": {
        "user Login": {
            "role": {
                "employee": [
                    {
                        "action": "Login",
                        "arguments": [
                            "username",
                            "password"
                        ]
                    }
                ]
            }
        }
    }
    }

    """
    Args:
        dom_structure: the scraped of dom tree of the webpage
        tech_doc: It the functioanlity document of the webpage

    Returns:
        DOM Elements that are related for doing the action defined in the tech_doc
    """
    prompt = f"""
            Analyze the given DOM structure of a webpage and the associated functionality document.
            Dom Structure:
            {dom_structure}

            Functionality document:
            {tech_doc}

            Your task is to determine if the elements and functionality described in the functionality document are present and correctly implemented in the DOM structure. Follow these steps:

            1. **Understand the Functionality Document**: Extract the key functionalities, actions, roles, and required arguments specified in the functionality document. Identify:
            - The actions. ("action")
            - The roles or users associated with these actions.
            - The required inputs or elements. ("arguments")

            2. **Parse the DOM Structure**: Break down the DOM structure to identify elements, attributes, and values that could correspond to the functionalities in the document. Focus on:
            - Input fields (`type`, `name`, `id`, `placeholder`, etc.).
            - Buttons (`type`, `class`, `role`, etc.).
            - Any associated elements or attributes relevant to the functionalities.

            3. **Verify Functionality Alignment**: Match the functionalities described in the document to the DOM structure. For each action:
            - Confirm that the required elements (e.g., input fields, buttons) exist.
            - Validate that attributes like `id`, `name`, `placeholder`, or `type` align with the functionality document's description.

            4. **Report Missing or Misaligned Elements**: Identify any functionality or elements that:
            - Are described in the functionality document but are missing in the DOM.
            - Exist in the DOM but do not match the functionality document's specifications.

            5. **Output Requirements**: Generate a dynamic, structured output that includes:
            - `"status"`: Whether all functionalities are present and correctly implemented (`true` or `false`).
            - `"missing_or_misaligned_elements"`: A list of missing or incorrectly implemented elements.
            - `"mapping"`: A clear mapping between the functionality document's specifications and the corresponding DOM elements.

            **Input Data Format:**
            1. **DOM Structure**: A JSON or textual representation of the webpage's DOM.
            2. **Functionality Document**: A JSON or textual description of the technical functionalities.

            **Example:**
            - If the functionality document specifies a "Login" button, check if the DOM has a `type="submit"` or a button element fulfilling that role.
            - If a "username" input field is required, confirm its presence and ensure its attributes match the description.

            Respond with a comprehensive analysis in JSON format.


    """
    messages = [
        {"role": "system", "content": "You are a helpful assistant."},
        {"role": "user", "content": prompt.strip()}
    ]
    response = openai.ChatCompletion.create(
    model="gpt-4o",
    messages=messages,
    max_tokens=1500,
    temperature=0
    )

    return response.choices[0].message['content'].strip()


def generate_selenium_script(element_dict = {}, testcases = []):
    testcases = ["Successful login using valid credentials", "Failed login using invalid credentials (random values for the negative test case)"]
    element_dict = {
        "status": True,
        "missing_or_misaligned_elements": [],
        "mapping": {
            "username": {
            "functionality_document": {
                "type": "text",
                "name": "username"
            },
            "dom_element": {
                "type": "text",
                "id": "username",
                "name": "username",
                "class": "oh-input w-100",
                "placeholder": "e.g. adam.luis@horilla.com"
            }
            },
            "password": {
            "functionality_document": {
                "type": "password",
                "name": "password"
            },
            "dom_element": {
                "type": "password",
                "id": "password",
                "name": "password",
                "class": "oh-input oh-input--password w-100",
                "placeholder": "Use alphanumeric characters"
            }
            },
            "login_button": {
            "functionality_document": {
                "action": "Login"
            },
            "dom_element": {
                "type": "submit",
                "class": "oh-btn oh-onboarding-card__button mt-4 oh-btn--secondary oh-btn--shadow w-100 mb-4",
                "role": "button"
            }
            }
        }
        }
    """
    Args:
        element_dict: The output of the comapring of the DOM structure and the technical document
        testcases: the list of test cases that that are generated with the help of the user stories

    Returns:
        Generates the pytest scripts using selenium for the required functioanlity and module 
    """
    prompt = f"""
            Based on the following data, generate Selenium test functions written in Python to be used with pytest for testing the login functionality of a webpage. The test functions should cover two scenarios: {" - ".join(testcases)}.

            Data:
            {element_dict}

            Notes:
            - generate only the selenium code
    """
    messages = [
        {"role": "system", "content": "You are a helpful assistant."},
        {"role": "user", "content": prompt.strip()}
    ]
    response = openai.ChatCompletion.create(
    model="gpt-4o",
    messages=messages,
    max_tokens=1500,
    temperature=0
    )

    return response.choices[0].message['content'].strip()


def parse_document(file_path):
    """
    Parse the input document to extract modules, user stories, roles, and actions.
    Returns:
        dict: A structured dictionary with extracted modules, user stories, roles, and actions.
    """
    structured_data = defaultdict(lambda: defaultdict(lambda: defaultdict(lambda: defaultdict(list))))

    doc = Document(file_path)
    document = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    # print(document)
    # Split the document into modules
    modules = re.split(r'\nModule Name: ', document)
    for module in modules:
        module = module.strip()
        if not module:
            continue

        # Extract the module name
        lines = module.splitlines()
        module_name = lines[0].strip()
        
        # Extract user stories and their details
        current_user_story = None
        current_role = None
        line_cn = 1
        while line_cn < len(lines):
            line = lines[line_cn]
        # for line in lines[1:]:
            line = line.strip()
            if line.startswith("User Story"):
                current_user_story = line.split(":", 1)[-1].strip()
                structured_data[module_name][current_user_story] = {
                    "roles": defaultdict(list)
                }
                line_cn += 1
            elif "Acceptance Criteria:" in line:
                line_cn += 1
                continue
            elif line.startswith("Role:"):
                current_role = line.split(":", 1)[-1].strip()
                # Initialize the role in the structured data if not already
                if current_user_story:
                    structured_data[module_name][current_user_story]["roles"][current_role] = []
                line_cn += 1
            elif any(action in line for action in ["Create", "Edit", "Update", "View", "Track", "Apply", "Schedule", "Store", "Define", "Assign", "Complete", "Approve", "Check-in/Check-out", "Cancel", "Upload", "Configure", "Request", "Download", "Raise"]):
                nw_line = line + " : " + lines[line_cn + 1]
                print(module_name, nw_line)
                action_match = re.match(r"(\w+ [^:]+): Arguments: (.+)", nw_line)
                if action_match:
                    action = action_match.group(1).strip()
                    arguments = [arg.strip() for arg in action_match.group(2).split(",")]
                    structured_data[module_name][current_user_story]["roles"][current_role].append({
                        "action": action,
                        "arguments": arguments
                    })
                line_cn += 2
                print(1, line_cn, len(lines))
            else:
                line_cn += 1
            if line_cn >= len(lines):
                break
    
    return json.dumps(structured_data, indent=4)

