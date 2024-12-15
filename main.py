import streamlit as st

from src.services.extract_page import extract_dom_attributes
from src.services.process import compare_dom_tech_doc, extract_text,generate_user_stories_and_criteria_specific, generate_selenium_script, parse_document
import json
from docx import Document


def main():
    st.title("BRD Automation System")

    st.sidebar.header("Inputs")
    brd_file = st.sidebar.file_uploader("Upload BRD Document", type=["pdf"])

    if st.sidebar.button("Generate Results"):
        if brd_file is None:
            st.sidebar.error("Please upload a BRD document before generating test cases.")
        else:
            st.subheader("Extracted BRD Requirements")
            text = extract_text(brd_file)

            st.subheader("Text Evaluation")
            with st.spinner("Executing test cases..."):
                # user_stories_and_criteria = generate_user_stories_and_criteria(text)
                user_stories_and_criteria_specific= generate_user_stories_and_criteria_specific(text)
                # st.write(user_stories_and_criteria_specific)

                doc = Document()
                output_file_path = "./src/services/result_file.docx"
                # # Add content to the document
                # if isinstance(user_stories_and_criteria_specific, list):
                #     for item in user_stories_and_criteria_specific:
                #         doc.add_paragraph(item)
                # else:
                #     doc.add_paragraph(user_stories_and_criteria_specific)

                # # Save the document to the specified path
                # doc.save(output_file_path)

                # creating the techincal document
                st.subheader("Created internal Technical Document")
                tech_doc = parse_document(output_file_path)
                st.write(tech_doc)

                ## creating the test cases
                # test_cases = generate_test_case()

                ## scraping the page for the complete dom tree
                page_url = "http://3.83.24.72:8000/"
                try:
                    # Extract and display DOM attributes
                    dom_data = extract_dom_attributes(page_url)
                    st.subheader("Extracted DOM of the page")
                    st.write(json.dump(dom_data, indent=4))
                except Exception as e:
                    print(f"An error occurred: {e}")

                st.subheader("Comapring the DOM")
                compare_data = compare_dom_tech_doc()
                st.write(compare_data)

                st.subheader("Generated Pytest scripts")
                sel_Script = generate_selenium_script()
                st.subheader("Scripts")
                st.write(sel_Script)

if __name__ == "__main__":
    main()
